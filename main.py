from db import init_db, get_all_items, add_rating
from docx import Document
from openpyxl import Workbook


def export_report(items):
    doc = Document()
    doc.add_heading('Итоговый отчёт: Виртуальный музей', 0)

    for _, type_, title, avg_rating, num_ratings, position in items:
        doc.add_paragraph(f"{position}. {type_} — {title}: {avg_rating:.2f} (оценок: {num_ratings})")

    doc.save('museum_summary_report.docx')

    wb = Workbook()
    ws = wb.active
    ws.title = "Отчёт"
    ws.append(["Позиция", "Тип", "Название", "Средняя оценка", "Кол-во оценок"])
    for _, type_, title, avg_rating, num_ratings, position in items:
        ws.append([position, type_, title, avg_rating, num_ratings])
    wb.save('museum_summary_report.xlsx')


def main():
    init_db()

    print("Добро пожаловать в Виртуальный музей!")
    print("Оцените объекты от 0 до 10.\n")

    while True:
        items = get_all_items()
        print("=" * 80)
        print(f"{'ID':<4} {'Тип':<10} {'Название':<40} {'Средняя':<10} {'Оценок':<8}")
        print("=" * 80)
        for item_id, type_, title, avg_rating, num_ratings, position in items:
            print(f"{item_id:<4} {type_:<10} {title:<40} {avg_rating:<10.2f} {num_ratings:<8}")
        print("=" * 80)

        choice = input("\nВведите ID объекта для оценки или 'стоп' для завершения: ").strip()
        if choice.lower() in ('стоп', 'stop', 'q', 'exit'):
            break

        if not choice.isdigit():
            print("Введите корректный ID (число).")
            continue

        item_id = int(choice)
        valid_ids = [row[0] for row in items]
        if item_id not in valid_ids:
            print("Такого ID нет. Попробуйте снова.")
            continue

        try:
            rating = float(input("Введите оценку (0–10): ").strip())
            if not (0 <= rating <= 10):
                print("Оценка должна быть от 0 до 10.")
                continue
        except ValueError:
            print("Введите число от 0 до 10.")
            continue

        add_rating(item_id, rating)
        print("Оценка сохранена!\n")

    print("\nИтоговая таблица:")
    print("=" * 80)
    print(f"{'№':<3} {'Тип':<10} {'Название':<30} {'Средняя':<10} {'Оценок':<8}")
    print("=" * 80)
    items = get_all_items()
    for item_id, type_, title, avg_rating, num_ratings, position in items:
        print(f"{position:<3} {type_:<10} {title:<30} {avg_rating:<10.2f} {num_ratings:<8}")
    print("=" * 80)

    export_report(items)
    print("\nОтчёт сохранён: museum_summary_report.docx, museum_summary_report.xlsx")


if __name__ == "__main__":
    main()